from docx import Document

from report_generator import (
    MANUAL_FIELD_KEYS,
    build_auto_manual_fields,
    build_entries,
    build_output_filename,
    build_template_context,
    classify_publication,
    extract_level,
    format_reference,
    render_report,
)


def sample_record(
    category: str = "Journal article",
    year: int = 2024,
    level: int | str = 2,
    code: str | None = None,
) -> dict:
    category_payload = {"name": {"en": category}}
    if code:
        category_payload["code"] = code

    return {
        "contributors": {"preview": [{"surname": "Doe", "first_name": "Jane"}]},
        "year_published": year,
        "title": {"en": "Sample Publication"},
        "journal": {"name": "Journal of Testing", "level": str(level)},
        "volume": "12",
        "pages": {"from": "10", "to": "20"},
        "links": [{"url_type": "DOI", "url": "https://doi.org/10.0000/example"}],
        "category": category_payload,
    }


def test_format_reference_contains_expected_fields():
    record = sample_record()
    reference = format_reference(record)
    assert "Doe, Jane" in reference
    assert "(2024)." in reference
    assert "Sample Publication" in reference
    assert "Journal of Testing" in reference
    assert "doi.org" in reference


def test_extract_level_prefers_known_fields():
    record = sample_record(level=1)
    assert extract_level(record) == "1"


def test_classify_publication_uses_category_and_level():
    record = sample_record(category="Monograph", level=2)
    assert classify_publication(record) == "publisert_monografi_niva2"

    record = sample_record(category="Anthology", level=1)
    assert classify_publication(record) == "publisert_antologi_niva1"

    record = sample_record(category="Book review", level=1)
    assert classify_publication(record) == "publisert_book_review"


def test_build_entries_filters_by_year():
    records = [sample_record(year=2023), sample_record(year=2024)]
    entries = build_entries(records, report_year=2024)
    assert len(entries) == 1
    assert entries[0].year == "2024"


def test_build_entries_skips_non_publication_categories():
    record = sample_record(category="Interview", code="MEDIAINTERVIEW")
    entries = build_entries([record], report_year=2024)
    assert entries == []


def test_build_template_context_populates_keys():
    records = [sample_record(category="Journal article", level=2)]
    entries = build_entries(records, report_year=2024)
    context = build_template_context(
        entries,
        report_year=2024,
        person_name="Jane Doe",
        institution_name="Test University",
        institution_name_secondary="",
        manual_fields={"formidling_kronikker": "Kronikk i Aftenposten."},
    )

    assert context["report_year"] == "2024"
    assert context["person_name"] == "Jane Doe"
    assert "publisert_artikkel_niva2" in context
    assert "Doe, Jane" in str(context["publisert_artikkel_niva2"])
    assert context["formidling_kronikker"] == "Kronikk i Aftenposten."
    for key in MANUAL_FIELD_KEYS:
        assert key in context


def test_build_auto_manual_fields_maps_programme_participation():
    record = {
        "category": {"code": "PROGRAMPARTICIP", "name": {"en": "Programme participation"}},
        "year_published": 2024,
        "title": {"en": "Psykt Interessant!"},
    }
    fields = build_auto_manual_fields([record], report_year=2024)
    assert "Psykt Interessant!" in fields["formidling_media"]


def test_build_output_filename_uses_person_name():
    filename = build_output_filename("Źse/Example", person_id=123, report_year=2024)
    assert filename.startswith("Aarsrapport_2024_")
    assert filename.endswith(".docx")
    assert "/" not in filename


def test_render_report_replaces_placeholders(tmp_path):
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "output.docx"
    doc = Document()
    doc.add_paragraph("{{ person_name }}")
    doc.save(template_path)

    render_report({"person_name": "Jane Doe"}, template_path, output_path)

    rendered = Document(output_path)
    text = "\n".join(p.text for p in rendered.paragraphs)
    assert "Jane Doe" in text
